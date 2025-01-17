from typing import List
import re
from langchain.docstore.document import Document
from langchain.embeddings.openai import OpenAIEmbeddings
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.vectorstores.faiss import FAISS
import faiss
import PyPDF2
import docx

def parse_text(text: str, filename: str) -> List[str]:
    text = re.sub(r"(\w+)-\n(\w+)", r"\1\2", text)
    text = re.sub(r"(?<!\n\s)\n(?!\s\n)", " ", text.strip())
    text = re.sub(r"\n\s*\n", "\n\n", text)
    return [text]

def parse_pdf(pdf_file) -> str:
    reader = PyPDF2.PdfReader(pdf_file)
    text = ""
    for page in reader.pages:
        text += page.extract_text() + "\n"
    return text

def parse_docx(docx_file) -> str:
    doc = docx.Document(docx_file)
    text = ""
    for para in doc.paragraphs:
        text += para.text + "\n"
    return text

def text_to_docs(text_list: List[str], filename: str) -> List[Document]:
    doc_chunks = []
    for text in text_list:
        text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=4000,
            separators=["\n\n", "\n", ".", "!", "?", ",", " ", ""],
            chunk_overlap=0,
        )
        chunks = text_splitter.split_text(text)
        for i, chunk in enumerate(chunks):
            doc = Document(
                page_content=chunk, metadata={"page": 1, "chunk": i, "filename": filename}
            )
            doc.metadata["source"] = f"{doc.metadata['page']}-{doc.metadata['chunk']}"
            doc_chunks.append(doc)
    return doc_chunks

def docs_to_index(docs, openai_api_key):
    index = FAISS.from_documents(docs, OpenAIEmbeddings(openai_api_key=openai_api_key))
    return index

def get_index_for_text_files(text_files: List[str], text_names: List[str], openai_api_key: str):
    documents = []
    for text_file, text_name in zip(text_files, text_names):
        if text_file.type == "application/pdf":
            text_content = parse_pdf(text_file)
        elif text_file.type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
            text_content = parse_docx(text_file)
        else:  # Assuming it's a text file
            text_content = text_file.read().decode("utf-8")
        
        text_list = parse_text(text_content, text_name)
        documents += text_to_docs(text_list, text_name)
    index = docs_to_index(documents, openai_api_key)
    return index
